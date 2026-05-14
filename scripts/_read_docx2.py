import docx
from docx.oxml.ns import qn

doc = docx.Document(r'C:\Users\xiaoy\Desktop\人工智能大作业报告.docx')

print("=== 段落内容 ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] style='{p.style.name}' bold={any(r.bold for r in p.runs)}: {repr(p.text)}")

print(f"\n=== 表格数量: {len(doc.tables)} ===")
for ti, table in enumerate(doc.tables):
    print(f"\n表格{ti}: {len(table.rows)}行 x {len(table.columns)}列")
    for ri, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        print(f"  行{ri}: {row_data}")

print("\n=== 样式列表 ===")
styles_used = set()
for p in doc.paragraphs:
    if p.text.strip():
        styles_used.add(p.style.name)
print(styles_used)

# 看看页眉页脚
print("\n=== 节信息 ===")
for i, section in enumerate(doc.sections):
    print(f"节{i}: 页边距 top={section.top_margin}, bottom={section.bottom_margin}")
    print(f"  页脚: {section.footer}")
    print(f"  页眉: {section.header}")